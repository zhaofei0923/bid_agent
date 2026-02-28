"""DOCX parser — extract text from Word documents using python-docx."""

import logging
from dataclasses import dataclass
from pathlib import Path

logger = logging.getLogger(__name__)


@dataclass
class ParsedDocx:
    """Parsed DOCX document."""

    full_text: str
    paragraphs: list[str]
    headings: list[dict]  # [{level, text, index}]
    tables: list[list[list[str]]]  # [table[row[cell]]]
    metadata: dict


def parse_docx(file_path: str | Path) -> ParsedDocx:
    """Parse a DOCX file and extract text, headings, and tables.

    Args:
        file_path: Path to the .docx file.

    Returns:
        ParsedDocx with structured content.
    """
    try:
        from docx import Document
    except ImportError:
        logger.error("python-docx not installed. Install with: pip install python-docx")
        raise

    file_path = Path(file_path)
    if not file_path.exists():
        raise FileNotFoundError(f"DOCX not found: {file_path}")

    doc = Document(str(file_path))

    paragraphs: list[str] = []
    headings: list[dict] = []

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            paragraphs.append(text)

        if para.style and para.style.name.startswith("Heading"):
            try:
                level = int(para.style.name.replace("Heading ", ""))
            except ValueError:
                level = 0
            headings.append({"level": level, "text": text, "index": i})

    # Extract tables
    tables: list[list[list[str]]] = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        tables.append(rows)

    full_text = "\n".join(paragraphs)

    # Core properties
    core = doc.core_properties
    metadata = {
        "title": core.title or "",
        "author": core.author or "",
        "subject": core.subject or "",
        "created": str(core.created) if core.created else "",
        "modified": str(core.modified) if core.modified else "",
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
    }

    return ParsedDocx(
        full_text=full_text,
        paragraphs=paragraphs,
        headings=headings,
        tables=tables,
        metadata=metadata,
    )


async def parse_docx_async(file_path: str | Path) -> ParsedDocx:
    """Async wrapper around parse_docx."""
    import asyncio

    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, parse_docx, file_path)
